
import streamlit as st
import numpy as np
import pandas as pd
import joblib
from io import BytesIO

try:
    from docx import Document
except Exception:
    Document = None

st.set_page_config(page_title='RAP Concrete RCPT & DME Predictor', layout='centered')

@st.cache_resource
def load_artifacts():
    feature_cols = joblib.load('feature_cols.pkl')
    ranges_df = pd.read_csv('input_feature_ranges.csv', index_col=0)
    rcpt_model = joblib.load('xgb_optuna_rcpt.pkl')
    params = joblib.load('xgb_optuna_params.pkl')
    return feature_cols, ranges_df, rcpt_model, params

feature_cols, ranges_df, rcpt_model, params = load_artifacts()

st.title('RAP Concrete Predictor')
st.caption('XGBoost regressor tuned with Optuna. Outputs: RCPT and DME (if model provided).')

st.markdown('#### Inputs')

col1, col2 = st.columns(2)

with col1:
    C = st.number_input('C (kg/m³)', value=float(ranges_df.loc['C','mean']), min_value=float(ranges_df.loc['C','min']), max_value=float(ranges_df.loc['C','max']))
    FA = st.number_input('FA (kg/m³)', value=float(ranges_df.loc['FA','mean']), min_value=float(ranges_df.loc['FA','min']), max_value=float(ranges_df.loc['FA','max']))
    FA_Binder = st.number_input('FA/Binder', value=float(ranges_df.loc['FA/Binder','mean']), min_value=float(ranges_df.loc['FA/Binder','min']), max_value=float(ranges_df.loc['FA/Binder','max']))
    A = st.number_input('A (kg/m³)', value=float(ranges_df.loc['A','mean']), min_value=float(ranges_df.loc['A','min']), max_value=float(ranges_df.loc['A','max']))

with col2:
    T_CRAP = st.number_input('T CRAP (kg/m³ or % as in training)', value=float(ranges_df.loc['T CRAP','mean']), min_value=float(ranges_df.loc['T CRAP','min']), max_value=float(ranges_df.loc['T CRAP','max']))
    CA = st.number_input('CA', value=float(ranges_df.loc['CA','mean']), min_value=float(ranges_df.loc['CA','min']), max_value=float(ranges_df.loc['CA','max']))
    Age = st.number_input('Age (days)', value=float(ranges_df.loc['Age','mean']), min_value=float(ranges_df.loc['Age','min']), max_value=float(ranges_df.loc['Age','max']))

input_dict = {
    'C': C,
    'FA': FA,
    'FA/Binder': FA_Binder,
    'A': A,
    'T CRAP': T_CRAP,
    'CA': CA,
    'Age': Age
}

X_input = pd.DataFrame([input_dict], columns=feature_cols)

st.markdown('#### Model details')
with st.expander('Show Optuna best hyperparameters'):
    st.write(params)

st.markdown('#### Prediction')

have_dme = os.path.exists('xgb_optuna_dme.pkl')
if have_dme:
    dme_model = joblib.load('xgb_optuna_dme.pkl')

predict_clicked = st.button('Predict RCPT and DME')

if predict_clicked:
    rcpt_pred = float(rcpt_model.predict(X_input)[0])
    dme_pred = None
    if have_dme:
        dme_pred = float(dme_model.predict(X_input)[0])

    m1, m2 = st.columns(2)
    with m1:
        st.metric('RCPT (Coulombs)', value=('{:,.2f}'.format(rcpt_pred)))
    with m2:
        if dme_pred is None:
            st.metric('DME', value='Model not found')
        else:
            st.metric('DME', value=('{:,.4f}'.format(dme_pred)))

    results_df = X_input.copy()
    results_df['RCPT_pred'] = rcpt_pred
    if dme_pred is not None:
        results_df['DME_pred'] = dme_pred

    st.markdown('#### Preview')
    st.dataframe(results_df, use_container_width=True)

    st.markdown('#### Download')
    csv_bytes = results_df.to_csv(index=False).encode('utf-8')
    st.download_button('Download results as CSV', data=csv_bytes, file_name='rap_predictions.csv', mime='text/csv')

    if Document is None:
        st.info('Install python-docx to enable DOCX download.')
    else:
        doc = Document()
        doc.add_heading('RAP Concrete Prediction Report', level=1)
        doc.add_paragraph('Model: XGBoost (Optuna-tuned)')
        doc.add_heading('Inputs', level=2)
        for k, v in input_dict.items():
            doc.add_paragraph(str(k) + ': ' + str(v))
        doc.add_heading('Predictions', level=2)
        doc.add_paragraph('RCPT (Coulombs): ' + ('{:,.2f}'.format(rcpt_pred)))
        if dme_pred is None:
            doc.add_paragraph('DME: model not available in deployment package')
        else:
            doc.add_paragraph('DME: ' + ('{:,.4f}'.format(dme_pred)))

        doc_buf = BytesIO()
        doc.save(doc_buf)
        st.download_button('Download report as DOCX', data=doc_buf.getvalue(), file_name='rap_prediction_report.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

st.markdown('---')
st.caption('Note: Input limits are based on the min/max values observed in the training dataset you provided.')
